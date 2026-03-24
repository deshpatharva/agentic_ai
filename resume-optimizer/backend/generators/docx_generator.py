"""
DOCX Generator
Creates a formatted .docx resume from plain resume text.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SECTION_HEADERS = re.compile(
    r"^(summary|professional\s+summary|objective|profile|about\s+me|"
    r"experience|work\s+experience|employment|work\s+history|professional\s+experience|"
    r"education|academic\s+background|qualifications|"
    r"skills|technical\s+skills|core\s+competencies|competencies|technologies|tools|"
    r"certifications?|licenses?|credentials?|"
    r"projects?|key\s+projects?|notable\s+projects?)$",
    re.IGNORECASE,
)

BULLET_PATTERN = re.compile(r"^[-•*\u2022\u2023\u25e6]\s+(.+)")

# Matches lines that contain a date range anywhere: "May 2022 – Aug 2022", "Dec 2023 – Current", etc.
DATE_PATTERN = re.compile(
    r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|"
    r"June|July|August|September|October|November|December)\s+\d{4}"
    r"\s*[-–—]\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|"
    r"March|April|June|July|August|September|October|November|December|\d{4}|Current|Present|Now)",
    re.IGNORECASE,
)

# Also match simple year ranges like "2020 – 2022"
YEAR_RANGE_PATTERN = re.compile(r"\b(20\d{2}|19\d{2})\s*[-–—]\s*(20\d{2}|19\d{2}|Current|Present)\b", re.IGNORECASE)

# Parenthetical annotations added by AI that should be stripped
AI_ANNOTATIONS = re.compile(
    r"\s*\((Established Company|Fortune 500|Startup|MNC|Listed Company|Public Company|Private Company)\)",
    re.IGNORECASE,
)


def _add_horizontal_rule(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _has_date(line: str) -> bool:
    return bool(DATE_PATTERN.search(line) or YEAR_RANGE_PATTERN.search(line))


def _split_company_date(line: str):
    """
    Split a company+date line into (company_part, date_part).
    Tries to find where the date starts and split there.
    """
    # Try month-year date first
    m = DATE_PATTERN.search(line)
    if not m:
        m = YEAR_RANGE_PATTERN.search(line)
    if m:
        company_part = line[:m.start()].strip().rstrip("–—-| \t")
        date_part = line[m.start():].strip()
        return company_part, date_part
    return line.strip(), ""


def _clean_ai_annotations(text: str) -> str:
    """Remove AI-added parenthetical company qualifiers."""
    return AI_ANNOTATIONS.sub("", text).strip()


def _is_name_line(line: str, seq_idx: int) -> bool:
    if seq_idx != 0:
        return False
    if SECTION_HEADERS.match(line.strip()):
        return False
    if BULLET_PATTERN.match(line.strip()):
        return False
    if len(line) > 60 or "@" in line or "http" in line.lower():
        return False
    return True


def _is_contact_line(line: str, seq_idx: int) -> bool:
    if seq_idx not in (1, 2, 3):
        return False
    indicators = ["@", "linkedin", "github", "http", "|", "·", "•", "(", "+"]
    return any(ind in line.lower() for ind in indicators)


def generate_docx(resume_text: str, output_path: str) -> str:
    """
    Generate a formatted .docx resume from plain text.

    Format rules:
    - Name: centered, large, bold, blue
    - Contact: centered, small, grey
    - Section headers: bold, blue, underlined
    - Company + date lines: company name BOLD, date right-aligned regular
    - Job title (line immediately after company+date): italic, not bold
    - Bullets: list style
    - Everything else: regular paragraph
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    lines = resume_text.splitlines()
    non_empty = [(i, ln) for i, ln in enumerate(lines) if ln.strip()]

    # Track whether the previous line was a company+date line
    prev_was_company = False

    for seq_idx, (orig_idx, line) in enumerate(non_empty):
        stripped = _clean_ai_annotations(line.strip())
        if not stripped:
            prev_was_company = False
            continue

        # ── Name ────────────────────────────────────────────────────────────
        if _is_name_line(stripped, seq_idx):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            prev_was_company = False
            continue

        # ── Contact info ────────────────────────────────────────────────────
        if _is_contact_line(stripped, seq_idx):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            prev_was_company = False
            continue

        # ── Section headers ─────────────────────────────────────────────────
        if SECTION_HEADERS.match(stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            _add_horizontal_rule(p)
            run = p.add_run(stripped.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            prev_was_company = False
            continue

        # ── Bullet points ───────────────────────────────────────────────────
        bullet_match = BULLET_PATTERN.match(stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(bullet_match.group(1))
            run.font.size = Pt(10.5)
            prev_was_company = False
            continue

        # ── Company + date line (e.g. "Hubot Inc, South Bend   Dec 2023 – Current") ──
        if _has_date(stripped):
            company_part, date_part = _split_company_date(stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(1)

            # Company name — bold
            run_company = p.add_run(company_part)
            run_company.bold = True
            run_company.font.size = Pt(10.5)

            # Spacer tab then date — regular, right-aligned via tab stop
            if date_part:
                run_tab = p.add_run("\t")
                run_date = p.add_run(date_part)
                run_date.bold = False
                run_date.font.size = Pt(10.5)
                run_date.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

                # Set a right tab stop at the right margin
                from docx.oxml import OxmlElement as OE
                pPr = p._p.get_or_add_pPr()
                tabs = OE("w:tabs")
                tab = OE("w:tab")
                tab.set(qn("w:val"), "right")
                tab.set(qn("w:pos"), "9360")  # ~6.5 inches in twips
                tabs.append(tab)
                pPr.append(tabs)

            prev_was_company = True
            continue

        # ── Job title (line immediately after company+date) ─────────────────
        if prev_was_company:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped)
            run.italic = True
            run.bold = False
            run.font.size = Pt(10.5)
            prev_was_company = False
            continue

        # ── Regular paragraph ───────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(stripped)
        run.font.size = Pt(10.5)
        prev_was_company = False

    doc.save(output_path)
    return output_path
