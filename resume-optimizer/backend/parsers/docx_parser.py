"""
DOCX Resume Parser
Extracts text from a .docx file and detects common resume sections.
"""

from docx import Document
from utils.section_parser import SECTION_PATTERNS


def _detect_sections(lines: list[str]) -> dict:
    """
    Walk through lines and assign each line to a section bucket.
    """
    sections: dict[str, list[str]] = {k: [] for k in SECTION_PATTERNS}
    sections["header"] = []

    current_section = "header"

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        matched = False
        for section_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(stripped):
                current_section = section_name
                matched = True
                break

        if not matched:
            sections[current_section].append(stripped)

    return {k: "\n".join(v) for k, v in sections.items() if v}


def parse_docx(file_path) -> dict:
    """
    Parse a .docx resume file.

    Args:
        file_path: Absolute path to the .docx file.

    Returns:
        {
            "raw_text": str,
            "sections": {
                "header": str,
                "summary": str,
                "experience": str,
                "education": str,
                "skills": str,
                ...
            }
        }
    """
    doc = Document(file_path)
    lines: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    # Also extract text from tables (e.g., two-column resume layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.extend(cell_text.splitlines())

    raw_text = "\n".join(lines)
    sections = _detect_sections(lines)

    return {
        "raw_text": raw_text,
        "sections": sections,
    }
