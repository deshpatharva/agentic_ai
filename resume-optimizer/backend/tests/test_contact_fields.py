"""Unit tests for profile contact fields: schema, text assembly, docx render."""


def test_sections_data_defaults_contact():
    from profiles.schemas import SectionsData
    s = SectionsData()
    assert s.contact.full_name == ""
    assert s.model_dump()["contact"]["email"] == ""


def test_sections_data_accepts_contact():
    from profiles.schemas import SectionsData
    s = SectionsData(contact={"full_name": "Jane Doe", "email": "jane@doe.com"})
    assert s.contact.full_name == "Jane Doe"
    assert s.contact.phone == ""  # unspecified fields default to empty


def test_sections_to_text_emits_name_and_contact_first():
    from utils.profile_utils import sections_to_text
    text = sections_to_text({
        "contact": {
            "full_name": "Jane Doe",
            "location": "Austin, TX",
            "email": "jane@doe.com",
            "phone": "5551234567",
            "linkedin": "linkedin.com/in/janedoe",
            "website": "",
        },
        "summary": "Engineer.",
    })
    lines = text.splitlines()
    assert lines[0] == "Jane Doe"
    assert lines[1] == "Austin, TX • jane@doe.com • 5551234567 • linkedin.com/in/janedoe"
    assert "Professional Summary" in lines


def test_sections_to_text_without_contact_starts_at_summary():
    from utils.profile_utils import sections_to_text
    text = sections_to_text({"summary": "Engineer."})
    assert text.splitlines()[0] == "Professional Summary"


def test_sections_to_text_name_only_no_contact_line():
    from utils.profile_utils import sections_to_text
    text = sections_to_text({
        "contact": {"full_name": "Jane Doe"},
        "summary": "Engineer.",
    })
    lines = text.splitlines()
    assert lines[0] == "Jane Doe"
    assert lines[1] == ""  # blank separator, no empty bullet-joined line


def test_generate_docx_renders_centered_name_and_contact(tmp_path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from generators.docx_generator import generate_docx
    from utils.profile_utils import sections_to_text

    text = sections_to_text({
        "contact": {
            "full_name": "Jane Doe",
            "location": "Austin, TX",
            "email": "jane@doe.com",
            "phone": "5551234567",
            "linkedin": "",
            "website": "",
        },
        "summary": "Engineer with experience.",
    })
    out = tmp_path / "out.docx"
    generate_docx(text, str(out))
    paras = [p for p in Document(str(out)).paragraphs if p.text.strip()]

    assert paras[0].text == "Jane Doe"
    assert paras[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paras[0].runs[0].bold is True

    assert paras[1].text.startswith("Austin, TX")
    assert paras[1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_generate_docx_renders_location_only_contact_centered(tmp_path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from generators.docx_generator import generate_docx
    from utils.profile_utils import sections_to_text

    text = sections_to_text({
        "contact": {"full_name": "Jane Doe", "location": "Austin, TX"},
        "summary": "Engineer.",
    })
    out = tmp_path / "out.docx"
    generate_docx(text, str(out))
    paras = [p for p in Document(str(out)).paragraphs if p.text.strip()]

    assert paras[0].text == "Jane Doe"
    assert paras[1].text == "Austin, TX"
    assert paras[1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_generate_docx_name_without_contact_keeps_header_styling(tmp_path):
    from docx import Document
    from generators.docx_generator import generate_docx
    from utils.profile_utils import sections_to_text

    text = sections_to_text({
        "contact": {"full_name": "Jane Doe"},
        "summary": "Engineer.",
    })
    out = tmp_path / "out.docx"
    generate_docx(text, str(out))
    paras = [p for p in Document(str(out)).paragraphs if p.text.strip()]

    assert paras[0].text == "Jane Doe"
    # The line after the name is the summary section header — it must keep
    # header styling (bold, uppercased), not be mistaken for contact info.
    assert paras[1].text == "PROFESSIONAL SUMMARY"
    assert paras[1].runs[0].bold is True


def test_generate_docx_date_line_after_name_not_styled_as_contact(tmp_path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from generators.docx_generator import generate_docx

    # Raw-upload text: no contact info, a company+date line directly after the name
    text = "John Smith\nAcme Corp  Jan 2020 – Dec 2023\n• Built things."
    out = tmp_path / "out.docx"
    generate_docx(text, str(out))
    paras = [p for p in Document(str(out)).paragraphs if p.text.strip()]

    assert paras[0].text == "John Smith"
    # Company+date keeps bold-company styling, not centered contact styling
    assert paras[1].alignment != WD_ALIGN_PARAGRAPH.CENTER
    assert paras[1].runs[0].bold is True


# ── Regression: contact/name heuristics (docx N4 / N7) ───────────────────────

def test_contact_heuristic_ignores_bare_5digit_in_prose():
    """A summary line with a 5-digit number must not be mistaken for a ZIP/contact."""
    from generators.docx_generator import _is_contact_line
    assert not _is_contact_line("Reached 50000 users; migrated 12000 records", 2)
    # Real ZIP formats are still recognized.
    assert _is_contact_line("Austin, TX 78701", 1)
    assert _is_contact_line("San Francisco, CA 94105-1234", 1)


def test_name_heuristic_rejects_non_name_first_lines():
    """A resume opening with an address / '5+ Years' / phone isn't the big name."""
    from generators.docx_generator import _is_name_line
    assert not _is_name_line("5+ Years of Experience in Software", 0)
    assert not _is_name_line("123 Main Street, Boston", 0)
    assert not _is_name_line("+1 (555) 123-4567", 0)
    # Genuine names still qualify.
    assert _is_name_line("Jane Doe", 0)
    assert _is_name_line("John Q. Smith", 0)


def test_name_line_with_digits_renders_as_contact_not_name(tmp_path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from generators.docx_generator import generate_docx

    # First line is an address (has digits) — must not become the 18pt blue name.
    text = "123 Main Street, Austin, TX 78701\nSummary\nEngineer."
    out = tmp_path / "out.docx"
    generate_docx(text, str(out))
    paras = [p for p in Document(str(out)).paragraphs if p.text.strip()]
    # Rendered as centered contact info, not an 18pt name run.
    assert paras[0].runs[0].font.size is None or paras[0].runs[0].font.size.pt < 18
